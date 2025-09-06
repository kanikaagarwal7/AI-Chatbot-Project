import fitz
import os
import uuid
import gridfs
import re                 # <-- added for keyword search
from bson import ObjectId
from pymongo import MongoClient
from datetime import datetime, timezone
from dotenv import load_dotenv
from ai21 import AI21Client
from ai21.models.chat import ChatMessage
from docx import Document   # for DOCX text extraction

# Load environment variables
load_dotenv()

# Get API key and MongoDB URI
api_key = os.getenv('AI21_API_KEY')
mongo_client = MongoClient("mongodb://localhost:27017")  # Update if using Atlas
db = mongo_client["chat_history_db"]
fs = gridfs.GridFS(db)
session_collection = db["chat_sessions"]

doc_content = ""   # will hold extracted text


# ------------------------------
# Helper: Highlight keywords
# ------------------------------
def highlight(text, keyword):
    """Highlight keyword in text with **word**"""
    pattern = re.compile(re.escape(keyword), re.IGNORECASE)
    return pattern.sub(lambda m: f"**{m.group(0)}**", text)


# ------------------------------
# Helper: Search inside embedded documents
# ------------------------------
def search_documents(query, doc_text):
    print(f"\n🔍 Search Results for '{query}' in documents:")
    matches = []
    for line in doc_text.split("\n"):
        if query.lower() in line.lower():
            matches.append(line.strip())

    if not matches:
        print("❌ No matches found in documents.")
    else:
        for snippet in matches:
            print(highlight(snippet, query))

#Helper: Search inside chat history (MongoDB)
def search_chat_history(query, session_id):
    print(f"\n💬 Search Results in Chat History for '{query}':")
    results = session_collection.find(
        {"_id": session_id, "chat_history": {"$exists": True}},
        {"chat_history": 1}
    )

    found = False
    for r in results:
        for chat in r.get("chat_history", []):
            if query.lower() in chat["question"].lower() or query.lower() in chat["answer"].lower():
                q = highlight(chat["question"], query)
                a = highlight(chat["answer"], query)
                print(f"Q: {q}")
                print(f"A: {a}\n---")
                found = True
    if not found:
        print("❌ No matches found in chat history.")


# ------------------------------
# Create or Continue a Session
# ------------------------------
session_choice = input("Do you want to (1) Create a new session or (2) Use existing? : ").strip()

# New session
if session_choice == "1":   
    session_id = str(uuid.uuid4())
    session_description = input("Enter a short description for this session: ").strip()
    session_collection.insert_one({
        "_id": session_id,
        "description": session_description,
        "created_at": datetime.now(timezone.utc),
        "documents": [],
        "chat_history": []
    })
    print(f"✅ New session created with ID: {session_id}")

    # Upload documents
    upload_choice = input("Do you want to upload documents? (yes/no): ").strip().lower()
    doc_content_parts = []

    if upload_choice == "yes":
        documents_uploaded = []

        # TXT
        try:
            with open("cat.txt", "rb") as txt_file:
                txt_file_id = fs.put(txt_file, filename="cat.txt")
                txt_file.seek(0)
                txt_content = txt_file.read().decode("utf-8")
                documents_uploaded.append({
                    "filename": "cat.txt",
                    "gridfs_id": str(txt_file_id),
                    "type": "txt"
                })
                doc_content_parts.append(txt_content)
                print(f"✅ Uploaded cat.txt with ID {txt_file_id}")
        except FileNotFoundError:
            pass

        # PDF
        pdf_path = "4thsemcorrected.pdf"
        if os.path.exists(pdf_path):
            with open(pdf_path, "rb") as pdf_file:
                pdf_file_id = fs.put(pdf_file, filename="4thsemcorrected.pdf")
                documents_uploaded.append({
                    "filename": "4thsemcorrected.pdf",
                    "gridfs_id": str(pdf_file_id),
                    "type": "pdf"
                })
                print(f"✅ Uploaded PDF with ID {pdf_file_id}")

            pdf = fitz.open(pdf_path)
            pdf_text = ""
            for page in pdf:
                pdf_text += page.get_text()
            pdf.close()
            doc_content_parts.append(pdf_text)

        # DOCX
        docx_path = "Bff.docx"
        if os.path.exists(docx_path):
            with open(docx_path, "rb") as docx_file:
                docx_file_id = fs.put(docx_file, filename="Bff.docx")
                documents_uploaded.append({
                    "filename": "Bff.docx",
                    "gridfs_id": str(docx_file_id),
                    "type": "docx"
                })
                print(f"✅ Uploaded Bff.docx with ID {docx_file_id}")

            try:
                doc = Document(docx_path)
                docx_text = "\n".join([p.text for p in doc.paragraphs])
                doc_content_parts.append(docx_text)
            except Exception as e:
                print(f"⚠️ Error reading Bff.docx: {e}")

        # Update session
        if documents_uploaded:
            session_collection.update_one(
                {"_id": session_id},
                {"$push": {"documents": {"$each": documents_uploaded}}}
            )

    doc_content = "\n\n".join(doc_content_parts)

# Use existing session
elif session_choice == "2":  
    print("\n📂 Available Sessions:")
    sessions = list(session_collection.find({}))
    if not sessions:
        print("❌ No existing sessions found. Exiting.")
        exit()

    for i, s in enumerate(sessions, 1):
        print(f"{i}. ID: {s['_id']} | Description: {s.get('description','No description')} | Created: {s['created_at']}")

    try:
        choice = int(input("\nEnter the number of the session you want to continue: ").strip())
        if 1 <= choice <= len(sessions):
            session = sessions[choice - 1]
            session_id = session['_id']
            print(f"ℹ️ Continuing session {session_id}")
             # Load already uploaded documents
            doc_content_parts = []
            print("\n📄 Documents already uploaded in this session:")

            seen = set()         # prevent duplicates
            for d in session.get("documents", []):
                key = (d.get("filename"), d.get("type"))
                if key in seen:
                    continue
                seen.add(key)

                try:
                    gridout = fs.get(ObjectId(d["gridfs_id"]))
                    content = gridout.read()

                    if d["type"] == "txt":
                        text = content.decode("utf-8")
                        doc_content_parts.append(text)

                    elif d["type"] == "pdf":
                        with open("temp.pdf", "wb") as f:
                            f.write(content)
                        pdf = fitz.open("temp.pdf")
                        pdf_text = ""
                        for page in pdf:
                            pdf_text += page.get_text()
                        pdf.close()
                        doc_content_parts.append(pdf_text)

                    elif d["type"] == "docx":
                        with open("temp.docx", "wb") as f:
                            f.write(content)
                        docx = Document("temp.docx")
                        docx_text = "\n".join([p.text for p in docx.paragraphs])
                        doc_content_parts.append(docx_text)

                    print(f"- {d.get('filename','Unknown')} (Type: {d.get('type')})")

                except Exception:
                    print(f"- {d.get('filename','Unknown')} (Type: {d.get('type')}) [content unavailable]")

            doc_content = "\n\n".join(doc_content_parts)

        else:
            print("❌ Invalid choice. Exiting.")
            exit()
    except ValueError:
        print("❌ Invalid input. Exiting.")
        exit()
else:
    print("❌ Invalid choice. Exiting.")
    exit()


# ------------------------------
# Chatbot Q&A Loop
# ------------------------------
client = AI21Client(api_key=api_key)

while True:
    action = input("\nWhat do you want to do? (ask/search_doc/search_chat/exit): ").strip().lower()
    if action == "exit":
        print("👋 Chat ended.")
        break

    elif action == "ask":
        mode = input("Choose mode: (1) Local documents or (2) Global knowledge: ").strip()
        if mode == "1":
            system = (
                "You are an assistant that must only answer using the following document. "
                "Do not use any external knowledge.\n\n"
                f"{doc_content}\n\n"
                "Instructions:\n"
                "- If the answer is found, respond with '(From local source)' followed by the answer.\n"
                "- If not found, respond with exactly: 'Not available in the document.'"
            )
        elif mode == "2":
            system = (
                "You are an AI assistant that answers using general knowledge.\n"
                "Important - Along with the answer, add this phrase: (From Global source)"
            )
        else:
            print("❌ Invalid mode.")
            continue

        user_input = input("Enter your question: ").strip()
        if user_input.lower() == "exit":
            break

        messages = [
            ChatMessage(content=system, role="system"),
            ChatMessage(content=user_input, role="user"),
        ]
        chat_completions = client.chat.completions.create(
            messages=messages,
            model="jamba-large",
        )
        response = chat_completions.choices[0].message.content
        print("\n🧠 AI Response:", response)

        session_collection.update_one(
            {"_id": session_id},
            {"$push": {
                "chat_history": {
                    "question": user_input,
                    "answer": response,
                    "mode": "local" if mode == "1" else "global",
                    "timestamp": datetime.now(timezone.utc)
                }
            }}
        )

    elif action == "search_doc":
        query = input("Enter keyword to search in uploaded documents: ").strip()
        search_documents(query, doc_content)

    elif action == "search_chat":
        query = input("Enter keyword to search in chat history: ").strip()
        search_chat_history(query, session_id)

    else:
        print("❌ Invalid action.")