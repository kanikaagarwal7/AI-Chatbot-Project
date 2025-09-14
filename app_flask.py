import os
import uuid
import fitz
import gridfs
import re
import base64
from bson import ObjectId
from flask import Flask, request, jsonify, render_template
from pymongo import MongoClient
from datetime import datetime, timezone
from dotenv import load_dotenv
from ai21 import AI21Client
from ai21.models.chat import ChatMessage
from docx import Document

# ------------------------------
# Load environment & setup
# ------------------------------
load_dotenv()
api_key = os.getenv('AI21_API_KEY')

# MongoDB setup
mongo_client = MongoClient("mongodb://localhost:27017")
db = mongo_client["chat_history_db"]
fs = gridfs.GridFS(db)
session_collection = db["chat_sessions"]

# AI21 setup
client = AI21Client(api_key=api_key)

# Flask app
app = Flask(__name__)


# ------------------------------
# Helpers
# ------------------------------
def highlight(text, keyword):
    """Highlight keyword in text with **word**"""
    pattern = re.compile(re.escape(keyword), re.IGNORECASE)
    return pattern.sub(lambda m: f"**{m.group(0)}**", text)


def extract_text_from_file(file_path, file_type):
    """Extract text depending on file type"""
    if file_type == "txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    elif file_type == "pdf":
        pdf = fitz.open(file_path)
        pdf_text = "".join([page.get_text() for page in pdf])
        pdf.close()
        return pdf_text
    elif file_type == "docx":
        doc = Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    return ""


# ------------------------------
# ROUTES
# ------------------------------

# 📌 Route 1: Create a new session
@app.route("/session/create", methods=["POST"])
def create_session():
    """Create a new chat session"""
    data = request.json
    session_id = str(uuid.uuid4())
    session_description = data.get("description", f"Session {session_id[:6]}")

    session_collection.insert_one({
        "_id": session_id,
        "description": session_description,
        "created_at": datetime.now(timezone.utc),
        "documents": [],
        "chat_history": [],
        "mode": "local"  # default mode
    })
    return jsonify({"session_id": session_id, "message": "Session created", "description": session_description}), 201


# 📌 Route 2: List all sessions
@app.route("/session/list", methods=["GET"])
def list_sessions():
    """List all saved chat sessions"""
    sessions = list(session_collection.find({}, {"chat_history": 0}))
    for s in sessions:
        s["_id"] = str(s["_id"])
    return jsonify(sessions)


# 📌 Route 3: Upload a document (txt/pdf/docx)
@app.route("/document/upload", methods=["POST"])
def upload_document():
    """
    Upload a document to a session.
    Supports:
    - multipart/form-data (file + session_id)
    - JSON body with base64 file content
    """
    session_id = None
    file_content = None
    filename = None
    file_type = None

    # Case 1: Multipart form-data
    if "file" in request.files and "session_id" in request.form:
        session_id = request.form.get("session_id")
        file = request.files["file"]

        filename = file.filename
        file_type = filename.split(".")[-1].lower()
        file_content = file.read()

    # Case 2: Raw JSON (base64 encoded file)
    elif request.is_json:
        data = request.get_json()
        session_id = data.get("session_id")
        filename = data.get("filename")
        file_type = filename.split(".")[-1].lower() if filename else None

        base64_data = data.get("file_content")
        if base64_data:
            try:
                file_content = base64.b64decode(base64_data)
            except Exception as e:
                return jsonify({"error": f"Invalid base64 data: {str(e)}"}), 400

    else:
        return jsonify({
            "error": "Invalid request format. Use form-data or JSON."
        }), 400

    # Final validation
    if not session_id or not file_content or not filename:
        return jsonify({
            "error": "Missing session_id, filename, or file content",
            "received_form": request.form.to_dict(),
            "received_files": list(request.files.keys())
        }), 400

    # Store in GridFS + Mongo
    file_id = fs.put(file_content, filename=filename)
    session_collection.update_one(
        {"_id": session_id},
        {"$push": {
            "documents": {
                "filename": filename,
                "gridfs_id": str(file_id),
                "type": file_type,
                "uploaded_at": datetime.now(timezone.utc)
            }
        }}
    )

    return jsonify({
        "message": f"✅ {filename} uploaded successfully",
        "session_id": session_id,
        "gridfs_id": str(file_id)
    })


# 📌 Route 4: Ask a question (local docs or global knowledge)
@app.route("/ask", methods=["POST"])
def ask_question():
    """Ask a question either using local docs or general knowledge"""
    data = request.json
    session_id = data.get("session_id")
    question = data.get("question")
    mode = data.get("mode")  # frontend will send mode

    session = session_collection.find_one({"_id": session_id})
    if not session:
        return jsonify({"error": "Session not found"}), 404

    # Collect text from documents
    doc_content_parts = []
    for d in session.get("documents", []):
        try:
            gridout = fs.get(ObjectId(d["gridfs_id"]))
            content = gridout.read()

            if d["type"] == "txt":
                doc_content_parts.append(content.decode("utf-8"))
            elif d["type"] == "pdf":
                with open("temp.pdf", "wb") as f:
                    f.write(content)
                pdf = fitz.open("temp.pdf")
                pdf_text = "".join([page.get_text() for page in pdf])
                pdf.close()
                doc_content_parts.append(pdf_text)
            elif d["type"] == "docx":
                with open("temp.docx", "wb") as f:
                    f.write(content)
                docx = Document("temp.docx")
                docx_text = "\n".join([p.text for p in docx.paragraphs])
                doc_content_parts.append(docx_text)
        except:
            continue

    doc_content = "\n\n".join(doc_content_parts)

    # Define system prompt
    if mode == "local":
        system = (
            "You are an assistant that must only answer using the following document. "
            "Do not use any external knowledge.\n\n"
            f"{doc_content}\n\n"
            "Instructions:\n"
            "- If the answer is found, respond with '(From local source)' followed by the answer.\n"
            "- If not found, respond with exactly: 'Not available in the document.'"
        )
    else:
        system = (
            "You are an AI assistant that answers using general knowledge.\n"
            "Important - Along with the answer, add this phrase: (From Global source)"
        )

    messages = [
        ChatMessage(content=system, role="system"),
        ChatMessage(content=question, role="user"),
    ]
    chat_completions = client.chat.completions.create(
        messages=messages,
        model="jamba-large",
    )
    response = chat_completions.choices[0].message.content

    # Save in chat history
    session_collection.update_one(
        {"_id": session_id},
        {"$push": {
            "chat_history": {
                "question": question,
                "answer": response,
                "mode": mode,
                "timestamp": datetime.now(timezone.utc)
            }
        }}
    )

    return jsonify({"answer": response})


# 📌 Route 5: Get Full Chat History
@app.route("/chat/history", methods=["POST"])
def get_chat_history():
    """Get all chat history for a given session_id"""
    data = request.json
    session_id = data.get("session_id")

    if not session_id:
        return jsonify({"error": "Missing session_id"}), 400

    session = session_collection.find_one({"_id": session_id})
    if not session:
        return jsonify({"error": "Session not found"}), 404

    history = session.get("chat_history", [])
    return jsonify({
        "session_id": session_id,
        "description": session.get("description", "No description"),
        "chat_history": history
    })


# 📌 Route 6: Toggle Mode for a session
@app.route("/session/toggle_mode", methods=["POST"])
def toggle_mode():
    """Toggle session mode between 'local' and 'global'"""
    data = request.get_json()
    session_id = data.get("session_id")

    if not session_id:
        return jsonify({"success": False, "message": "Missing session_id"}), 400

    session = session_collection.find_one({"_id": session_id})
    if not session:
        return jsonify({"success": False, "message": "Session not found"}), 404

    current_mode = session.get("mode", "local")
    new_mode = "global" if current_mode == "local" else "local"

    session_collection.update_one(
        {"_id": session_id},
        {"$set": {"mode": new_mode}}
    )

    return jsonify({"success": True, "new_mode": new_mode})


# 📌 Route 7: Search inside uploaded documents
@app.route("/search/documents", methods=["POST"])
def search_documents_api():
    """Search for a keyword inside uploaded documents"""
    data = request.json
    session_id = data.get("session_id")
    query = data.get("q")

    session = session_collection.find_one({"_id": session_id})
    if not session:
        return jsonify({"error": "Session not found"}), 404

    doc_content_parts = []
    for d in session.get("documents", []):
        try:
            gridout = fs.get(ObjectId(d["gridfs_id"]))
            content = gridout.read()
            if d["type"] == "txt":
                doc_content_parts.append(content.decode("utf-8"))
            elif d["type"] == "pdf":
                with open("temp.pdf", "wb") as f:
                    f.write(content)
                pdf = fitz.open("temp.pdf")
                pdf_text = "".join([page.get_text() for page in pdf])
                pdf.close()
                doc_content_parts.append(pdf_text)
            elif d["type"] == "docx":
                with open("temp.docx", "wb") as f:
                    f.write(content)
                docx = Document("temp.docx")
                docx_text = "\n".join([p.text for p in docx.paragraphs])
                doc_content_parts.append(docx_text)
        except:
            continue

    matches = []
    for line in "\n".join(doc_content_parts).split("\n"):
        if query.lower() in line.lower():
            matches.append(highlight(line.strip(), query))

    return jsonify({"query": query, "matches": matches})


# 📌 Route 8: Search inside chat history
@app.route("/search/chat", methods=["POST"])
def search_chat_api():
    """Search for a keyword inside chat history"""
    data = request.json
    session_id = data.get("session_id")
    query = data.get("q")

    session = session_collection.find_one({"_id": session_id}, {"chat_history": 1})
    if not session:
        return jsonify({"error": "Session not found"}), 404

    matches = []
    for chat in session.get("chat_history", []):
        if query.lower() in chat["question"].lower() or query.lower() in chat["answer"].lower():
            matches.append({
                "question": highlight(chat["question"], query),
                "answer": highlight(chat["answer"], query)
            })

    return jsonify({"query": query, "matches": matches})


# 📌 Route 9: Delete a session
@app.route("/session/delete", methods=["POST"])
def delete_session():
    data = request.get_json()
    session_id = data.get("session_id")

    if not session_id:
        return jsonify({"success": False, "message": "Session ID missing"}), 400

    try:
        result = session_collection.delete_one({"_id": session_id})

        if result.deleted_count > 0:
            return jsonify({"success": True, "message": "Session deleted"})
        else:
            return jsonify({"success": False, "message": "Session not found"}), 404
    except Exception as e:
        return jsonify({"success": False, "message": str(e)}), 500
    
# 📌 Route 10: to list the documents in the session list
@app.route("/document/list", methods=["POST"])
def list_documents():
    data = request.json
    session_id = data.get("session_id")
    if not session_id:
        return jsonify({"error": "Missing session_id"}), 400

    session = session_collection.find_one({"_id": session_id})
    if not session:
        return jsonify({"error": "Session not found"}), 404

    # Assuming you store docs in session like: { "documents": [ { "filename": "x.pdf", "file_id": "..."} ] }
    docs = session.get("documents", [])
    return jsonify({"documents": docs})

# 📌 Route 11: to delete the documents from the session list
@app.route("/document/delete", methods=["POST"])
def delete_document():
    data = request.json
    session_id = data.get("session_id")
    filename = data.get("filename")

    if not session_id or not filename:
        return jsonify({"message": "Missing session_id or filename"}), 400

    session = session_collection.find_one({"_id": session_id})
    if not session:
        return jsonify({"message": "Session not found"}), 404

    documents = session.get("documents", [])
    updated_docs = [doc for doc in documents if doc.get("filename") != filename]

    if len(updated_docs) == len(documents):
        return jsonify({"message": "File not found in session"}), 404

    # ✅ Update session with remaining docs
    session_collection.update_one(
        {"_id": session_id}, {"$set": {"documents": updated_docs}}
    )

    # ✅ Optionally also delete from GridFS
    for doc in documents:
        if doc.get("filename") == filename:
            try:
                fs.delete(ObjectId(doc["gridfs_id"]))
            except Exception as e:
                print("GridFS delete error:", e)

    return jsonify({"message": f"{filename} deleted successfully"})

# Home route
@app.route("/", methods=["GET"])
def home():
    return render_template("index.html")


# ------------------------------
# Run Flask App
# ------------------------------
if __name__ == "__main__":
    app.run(debug=True, port=5000)
