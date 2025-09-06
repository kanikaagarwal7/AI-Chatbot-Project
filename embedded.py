import fitz
import os
import uuid
import gridfs
from bson import ObjectId
from pymongo import MongoClient
from datetime import datetime, timezone
from dotenv import load_dotenv
from ai21 import AI21Client
from ai21.models.chat import ChatMessage
from docx import Document   # <-- for DOCX text extraction

# Load environment variables
load_dotenv()

# Get API key and MongoDB URI
api_key = os.getenv('AI21_API_KEY')
mongo_client = MongoClient("mongodb://localhost:27017")  # Update if using Atlas
db = mongo_client["chat_history_db"]
fs = gridfs.GridFS(db)  
session_collection = db["chat_sessions"]

doc_content = ""   # will hold extracted text

# ------------------------------
# Create or Continue a Session
# ------------------------------
session_choice = input("Do you want to (1) Create a new session or (2) Use existing? : ").strip()

if session_choice == "1":   # New session
    session_id = str(uuid.uuid4())
    session_description = input("Enter a short description for this session: ").strip()
    session_collection.insert_one({
        "_id": session_id,
        "description": session_description,
        "created_at": datetime.now(timezone.utc),
        "documents": [],
        "chat_history": []
    })
    print(f"✅ New session created with ID: {session_id}")

    # ------------------------------
    # Upload Documents with Metadata
    # ------------------------------
    upload_choice = input("Do you want to upload documents? (yes/no): ").strip().lower()
    doc_content_parts = []

    if upload_choice == "yes":
        documents_uploaded = []

        # TXT
        try:
            with open("cat.txt", "rb") as txt_file:
                txt_file_id = fs.put(txt_file, filename="cat.txt")
                txt_file.seek(0)
                txt_content = txt_file.read().decode("utf-8")
                documents_uploaded.append({
                    "filename": "cat.txt",
                    "gridfs_id": str(txt_file_id),
                    "type": "txt"
                })
                doc_content_parts.append(txt_content)
                print(f"✅ Uploaded cat.txt with ID {txt_file_id}")
        except FileNotFoundError:
            pass

        # PDF
        pdf_path = "4thsemcorrected.pdf"
        if os.path.exists(pdf_path):
            with open(pdf_path, "rb") as pdf_file:
                pdf_file_id = fs.put(pdf_file, filename="4thsemcorrected.pdf")
                documents_uploaded.append({
                    "filename": "4thsemcorrected.pdf",
                    "gridfs_id": str(pdf_file_id),
                    "type": "pdf"
                })
                print(f"✅ Uploaded PDF with ID {pdf_file_id}")

            pdf = fitz.open(pdf_path)
            pdf_text = ""
            for page in pdf:
                pdf_text += page.get_text()
            pdf.close()
            doc_content_parts.append(pdf_text)

        # DOCX
        docx_path = "Bff.docx"
        if os.path.exists(docx_path):
            with open(docx_path, "rb") as docx_file:
                docx_file_id = fs.put(docx_file, filename="Bff.docx")
                documents_uploaded.append({
                    "filename": "Bff.docx",
                    "gridfs_id": str(docx_file_id),
                    "type": "docx"
                })
                print(f"✅ Uploaded Bff.docx with ID {docx_file_id}")

            try:
                doc = Document(docx_path)
                docx_text = "\n".join([p.text for p in doc.paragraphs])
                doc_content_parts.append(docx_text)
            except Exception as e:
                print(f"⚠️ Error reading Bff.docx: {e}")

        # Update session with embedded document info
        if documents_uploaded:
            session_collection.update_one(
                {"_id": session_id},
                {"$push": {"documents": {"$each": documents_uploaded}}}
            )

    # Combine all text
    doc_content = "\n\n".join(doc_content_parts)

elif session_choice == "2":  # Use existing session
    print("\n📂 Available Sessions:")
    sessions = list(session_collection.find({}))
    if not sessions:
        print("❌ No existing sessions found. Exiting.")
        exit()

    for i, s in enumerate(sessions, 1):
        print(f"{i}. ID: {s['_id']} | Description: {s.get('description','No description')} | Created: {s['created_at']}")

    try:
        choice = int(input("\nEnter the number of the session you want to continue: ").strip())
        if 1 <= choice <= len(sessions):
            session = sessions[choice - 1]
            session_id = session['_id']
            print(f"ℹ️ Continuing session {session_id}")

            # Load already uploaded documents
            doc_content_parts = []
            print("\n📄 Documents already uploaded in this session:")

            seen = set()  # prevent duplicates
            for d in session.get("documents", []):
                key = (d.get("filename"), d.get("type"))  # unique key
                if key in seen:
                    continue
                seen.add(key)

                try:
                    gridout = fs.get(ObjectId(d["gridfs_id"]))
                    content = gridout.read()

                    if d["type"] == "txt":
                        text = content.decode("utf-8")
                        doc_content_parts.append(text)

                    elif d["type"] == "pdf":
                        with open("temp.pdf", "wb") as f:
                            f.write(content)
                        pdf = fitz.open("temp.pdf")
                        pdf_text = ""
                        for page in pdf:
                            pdf_text += page.get_text()
                        pdf.close()
                        doc_content_parts.append(pdf_text)

                    elif d["type"] == "docx":
                        with open("temp.docx", "wb") as f:
                            f.write(content)
                        docx = Document("temp.docx")
                        docx_text = "\n".join([p.text for p in docx.paragraphs])
                        doc_content_parts.append(docx_text)

                    print(f"- {d.get('filename','Unknown')} (Type: {d.get('type')})")

                except Exception:
                    print(f"- {d.get('filename','Unknown')} (Type: {d.get('type')}) [content unavailable]")

            doc_content = "\n\n".join(doc_content_parts)

        else:
            print("❌ Invalid choice. Exiting.")
            exit()
    except ValueError:
        print("❌ Invalid input. Exiting.")
        exit()
else:
    print("❌ Invalid choice. Exiting.")
    exit()

# ------------------------------
# Chatbot Q&A Loop
# ------------------------------
client = AI21Client(api_key=api_key)

while True:
    ask_choice = input("\nDo you want to ask a question? (yes/exit): ").strip().lower()
    if ask_choice == "exit":
        print("👋 Chat ended.")
        break

    # Mode selection
    mode = input("Choose mode: (1) Local documents or (2) Global knowledge: ").strip()

    if mode == "1":
        system = (
            "You are an assistant that must only answer using the following document. "
            "Do not use any external knowledge.\n\n"
            f"{doc_content}\n\n"
            "Instructions:\n"
            "- If the answer is found, respond with '(From local source)' followed by the answer.\n"
            "- If not found, respond with exactly: 'Not available in the document.'"
        )
    elif mode == "2":
        system = (
            "You are an AI assistant that answers using general knowledge.\n"
            "Important - Along with the answer, add this phrase: (From Global source)"
        )
    else:
        print("❌ Invalid mode.")
        continue

    # Take user question
    user_input = input("Enter your question: ").strip()
    if user_input.lower() == "exit":
        break

    # Get AI response
    messages = [
        ChatMessage(content=system, role="system"),
        ChatMessage(content=user_input, role="user"),
    ]
    chat_completions = client.chat.completions.create(
        messages=messages,
        model="jamba-large",
    )
    response = chat_completions.choices[0].message.content
    print("\n🧠 AI Response:", response)

    # Store Q&A inside embedded chat_history
    session_collection.update_one(
        {"_id": session_id},
        {"$push": {
            "chat_history": {
                "question": user_input,
                "answer": response,
                "mode": "local" if mode == "1" else "global",
                "timestamp": datetime.now(timezone.utc)
            }
        }}
    )
